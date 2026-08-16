import tkinter as tk
from tkinter import ttk, messagebox, filedialog

from models import Process
from sjf_preemptive import run_sjf
from fcfs import run_fcfs
from comparator import calculate_average


class SchedulerApp:
    def __init__(self, root):
        self.root = root
        self.root.title("CPU Scheduling Simulator")
        self.root.geometry("1200x800")

        self.processes = []
        self.completed_processes = []
        self.gantt_data = []

        # =========================
        # TITLE
        # =========================
        title = tk.Label(
            self.root,
            text="CPU SCHEDULING SIMULATOR",
            font=("Arial", 22, "bold")
        )
        title.pack(pady=15)

        # =========================
        # PROCESS INPUT
        # =========================
        input_frame = ttk.LabelFrame(
            self.root,
            text="Process Input",
            padding=15
        )
        input_frame.pack(
            fill="x",
            padx=20,
            pady=5
        )

        ttk.Label(
            input_frame,
            text="PID:"
        ).grid(
            row=0,
            column=0,
            padx=5,
            pady=5
        )

        self.pid_entry = ttk.Entry(
            input_frame,
            width=15
        )
        self.pid_entry.grid(
            row=0,
            column=1,
            padx=5,
            pady=5
        )

        ttk.Label(
            input_frame,
            text="Arrival Time:"
        ).grid(
            row=0,
            column=2,
            padx=5,
            pady=5
        )

        self.arrival_entry = ttk.Entry(
            input_frame,
            width=15
        )
        self.arrival_entry.grid(
            row=0,
            column=3,
            padx=5,
            pady=5
        )

        ttk.Label(
            input_frame,
            text="Burst Time:"
        ).grid(
            row=0,
            column=4,
            padx=5,
            pady=5
        )

        self.burst_entry = ttk.Entry(
            input_frame,
            width=15
        )
        self.burst_entry.grid(
            row=0,
            column=5,
            padx=5,
            pady=5
        )

        add_button = ttk.Button(
            input_frame,
            text="Add Process",
            command=self.add_process
        )
        add_button.grid(
            row=0,
            column=6,
            padx=10,
            pady=5
        )

        # =========================
        # PROCESS TABLE - 7 COLUMNS
        # =========================
        table_frame = ttk.LabelFrame(
            self.root,
            text="Process Table",
            padding=10
        )
        table_frame.pack(
            fill="both",
            expand=True,
            padx=20,
            pady=5
        )

        columns = (
            "pid",
            "arrival",
            "burst",
            "completion",
            "turnaround",
            "waiting",
            "response"
        )

        self.tree = ttk.Treeview(
            table_frame,
            columns=columns,
            show="headings",
            height=8
        )

        self.tree.heading("pid", text="PID")
        self.tree.heading("arrival", text="AT")
        self.tree.heading("burst", text="BT")
        self.tree.heading("completion", text="CT")
        self.tree.heading("turnaround", text="TAT")
        self.tree.heading("waiting", text="WT")
        self.tree.heading("response", text="RT")

        for column in columns:
            self.tree.column(
                column,
                width=100,
                anchor="center"
            )

        self.tree.pack(
            fill="both",
            expand=True
        )

        # =========================
        # BUTTONS
        # =========================
        button_frame = ttk.Frame(self.root)
        button_frame.pack(pady=8)

        delete_button = ttk.Button(
            button_frame,
            text="Delete Selected",
            command=self.delete_selected
        )
        delete_button.grid(
            row=0,
            column=0,
            padx=5
        )

        reset_button = ttk.Button(
            button_frame,
            text="Reset",
            command=self.reset_all
        )
        reset_button.grid(
            row=0,
            column=1,
            padx=5
        )

        sjf_button = ttk.Button(
            button_frame,
            text="Run SJF Preemptive",
            command=self.run_sjf_algorithm
        )
        sjf_button.grid(
            row=0,
            column=2,
            padx=5
        ) 
        fcfs_button = ttk.Button(
            button_frame,
            text="Run FCFS",
            command=self.run_fcfs_algorithm
        )
        fcfs_button.grid(
            row=0,
            column=3,
            padx=5
        ) 
        compare_button = ttk.Button(
            button_frame,
            text="Compare",
            command=self.compare_algorithms
        )
        compare_button.grid(
            row=0,
            column=4,
            padx=5
        )

        export_button = ttk.Button(
            button_frame,
            text="Export DOCX",
            command=self.export_docx
        )
        export_button.grid(
            row=0,
            column=5,
            padx=5
        )

        # =========================
        # GANTT CHART
        # =========================
        gantt_frame = ttk.LabelFrame(
            self.root,
            text="Gantt Chart",
            padding=10
        )
        gantt_frame.pack(
            fill="x",
            padx=20,
            pady=5
        )

        self.canvas = tk.Canvas(
            gantt_frame,
            height=150,
            background="white"
        )
        self.canvas.pack(
            fill="x",
            expand=True
        )

    # =========================
    # ADD PROCESS
    # =========================
    def add_process(self):
        pid = self.pid_entry.get().strip()
        arrival_time = self.arrival_entry.get().strip()
        burst_time = self.burst_entry.get().strip()

        if not pid or not arrival_time or not burst_time:
            messagebox.showwarning(
                "Input Error",
                "Please enter all information."
            )
            return

        try:
            arrival_time = int(arrival_time)
            burst_time = int(burst_time)
        except ValueError:
            messagebox.showerror(
                "Input Error",
                "Arrival Time and Burst Time must be integers."
            )
            return

        if arrival_time < 0:
            messagebox.showerror(
                "Input Error",
                "Arrival Time must be greater than or equal to 0."
            )
            return

        if burst_time <= 0:
            messagebox.showerror(
                "Input Error",
                "Burst Time must be greater than 0."
            )
            return

        for process in self.processes:
            if process.pid == pid:
                messagebox.showerror(
                    "Input Error",
                    "PID already exists."
                )
                return

        process = Process(
            pid,
            arrival_time,
            burst_time
        )

        self.processes.append(process)

        # Khi input thay đổi, kết quả thuật toán cũ không còn hợp lệ.
        self.completed_processes = []
        self.gantt_data = []

        self.refresh_table(self.processes)
        self.canvas.delete("all")
        self.clear_entries()

    # =========================
    # CLEAR INPUT
    # =========================
    def clear_entries(self):
        self.pid_entry.delete(0, tk.END)
        self.arrival_entry.delete(0, tk.END)
        self.burst_entry.delete(0, tk.END)

        self.pid_entry.focus()

    # =========================
    # REFRESH TABLE
    # =========================
    def refresh_table(self, processes):
        for item in self.tree.get_children():
            self.tree.delete(item)

        for process in processes:
            self.tree.insert(
                "",
                tk.END,
                values=(
                    process.pid,
                    process.arrival_time,
                    process.burst_time,
                    process.completion_time,
                    process.turnaround_time,
                    process.waiting_time,
                    process.response_time
                )
            )

    # =========================
    # DELETE PROCESS
    # =========================
    def delete_selected(self):
        selected = self.tree.selection()

        if not selected:
            messagebox.showwarning(
                "Delete Process",
                "Please select a process."
            )
            return

        selected_pids = []

        for item in selected:
            values = self.tree.item(
                item,
                "values"
            )

            if values:
                selected_pids.append(
                    str(values[0])
                )

        self.processes = [
            process
            for process in self.processes
            if process.pid not in selected_pids
        ]

        self.completed_processes = []
        self.gantt_data = []

        self.refresh_table(self.processes)
        self.canvas.delete("all")

    # =========================
    # RESET
    # =========================
    def reset_all(self):
        self.processes.clear()
        self.completed_processes.clear()
        self.gantt_data.clear()

        self.refresh_table([])
        self.canvas.delete("all")
        self.clear_entries()

    # =========================
    # RUN SJF PREEMPTIVE
    # =========================
    def run_sjf_algorithm(self):
        if not self.processes:
            messagebox.showwarning(
                "SJF Preemptive",
                "Please add at least one process."
            )
            return

        try:
            completed_processes, gantt_chart = run_sjf(
                self.processes
            )
        except Exception as error:
            messagebox.showerror(
                "SJF Error",
                str(error)
            )
            return

        self.completed_processes = completed_processes
        self.gantt_data = gantt_chart

        # Hiển thị CT, TAT, WT, RT
        self.refresh_table(
            self.completed_processes
        )

        # Vẽ Gantt thật từ SRTF
        self.draw_gantt_chart(
            self.gantt_data
        )

        messagebox.showinfo(
            "SJF Preemptive",
            "SJF Preemptive completed successfully."
        )
            # =========================
    # RUN FCFS
    # =========================
    def run_fcfs_algorithm(self):
        if not self.processes:
            messagebox.showwarning(
                "FCFS",
                "Please add at least one process."
            )
            return

        try:
            completed_processes, gantt_chart = run_fcfs(
                self.processes
            )
        except Exception as error:
            messagebox.showerror(
                "FCFS Error",
                str(error)
            )
            return

        self.completed_processes = completed_processes
        self.gantt_data = gantt_chart

        self.refresh_table(
            self.completed_processes
        )

        self.draw_gantt_chart(
            self.gantt_data
        )

        messagebox.showinfo(
            "FCFS",
            "FCFS completed successfully."
        )
            # =========================
    # COMPARE SJF VS FCFS
    # =========================
    def compare_algorithms(self):
        if not self.processes:
            messagebox.showwarning(
                "Compare",
                "Please add at least one process."
            )
            return

        try:
            sjf_input = [
                Process(
                    process.pid,
                    process.arrival_time,
                    process.burst_time
                )
                for process in self.processes
            ]

            fcfs_input = [
                Process(
                    process.pid,
                    process.arrival_time,
                    process.burst_time
                )
                for process in self.processes
            ]

            sjf_result, _ = run_sjf(sjf_input)
            fcfs_result, _ = run_fcfs(fcfs_input)

            sjf_wt, sjf_tat = calculate_average(sjf_result)
            fcfs_wt, fcfs_tat = calculate_average(fcfs_result)

            if sjf_wt < fcfs_wt:
                wt_result = "SJF Preemptive has lower Average WT."
            elif fcfs_wt < sjf_wt:
                wt_result = "FCFS has lower Average WT."
            else:
                wt_result = "Both have the same Average WT."

            if sjf_tat < fcfs_tat:
                tat_result = "SJF Preemptive has lower Average TAT."
            elif fcfs_tat < sjf_tat:
                tat_result = "FCFS has lower Average TAT."
            else:
                tat_result = "Both have the same Average TAT."

            result_text = (
                "SJF PREEMPTIVE (SRTF)\n"
                f"Average WT: {sjf_wt:.2f}\n"
                f"Average TAT: {sjf_tat:.2f}\n\n"
                "FCFS\n"
                f"Average WT: {fcfs_wt:.2f}\n"
                f"Average TAT: {fcfs_tat:.2f}\n\n"
                "COMPARISON\n"
                f"{wt_result}\n"
                f"{tat_result}"
            )

            messagebox.showinfo(
                "SJF Preemptive vs FCFS",
                result_text
            )

        except Exception as error:
            messagebox.showerror(
                "Compare Error",
                str(error)
            )

    # =========================
    # DRAW REAL SJF GANTT CHART
    # =========================
    def draw_gantt_chart(self, gantt_chart=None):
        self.canvas.delete("all")

        if gantt_chart is None:
            gantt_chart = self.gantt_data

        if not gantt_chart:
            self.canvas.create_text(
                500,
                70,
                text="Run SJF Preemptive to generate Gantt Chart",
                font=("Arial", 13)
            )
            return

        first_start = gantt_chart[0][1]
        final_end = gantt_chart[-1][2]

        total_time = final_end - first_start

        if total_time <= 0:
            return

        canvas_width = self.canvas.winfo_width()

        if canvas_width <= 1:
            canvas_width = 1000

        left_margin = 40
        right_margin = 40

        available_width = (
            canvas_width
            - left_margin
            - right_margin
        )

        scale = available_width / total_time

        y = 30
        height = 55

        for pid, start_time, end_time in gantt_chart:
            x1 = (
                left_margin
                + (start_time - first_start) * scale
            )

            x2 = (
                left_margin
                + (end_time - first_start) * scale
            )

            self.canvas.create_rectangle(
                x1,
                y,
                x2,
                y + height,
                outline="black",
                width=2
            )

            self.canvas.create_text(
                (x1 + x2) / 2,
                y + height / 2,
                text=pid,
                font=("Arial", 11, "bold")
            )

            self.canvas.create_text(
                x1,
                y + height + 20,
                text=str(start_time)
            )

        self.canvas.create_text(
            left_margin + total_time * scale,
            y + height + 20,
            text=str(final_end)
        )

    # =========================
    # EXPORT DOCX
    # =========================
    def export_docx(self):
        if not self.processes:
            messagebox.showwarning(
                "Export DOCX",
                "There is no process data to export."
            )
            return

        if not self.completed_processes:
            messagebox.showwarning(
                "Export DOCX",
                "Run SJF Preemptive before exporting the report."
            )
            return

        try:
            from docx import Document
        except ImportError:
            messagebox.showerror(
                "Missing Library",
                "python-docx is not installed."
            )
            return

        file_path = filedialog.asksaveasfilename(
            title="Save SJF Preemptive Report",
            defaultextension=".docx",
            filetypes=[
                ("Word Document", "*.docx")
            ]
        )

        if not file_path:
            return

        document = Document()

        document.add_heading(
            "CPU Scheduling Simulator",
            level=1
        )

        document.add_heading(
            "SJF Preemptive (SRTF) Results",
            level=2
        )

        # 7-column result table
        table = document.add_table(
            rows=1,
            cols=7
        )

        table.style = "Table Grid"

        headers = [
            "PID",
            "AT",
            "BT",
            "CT",
            "TAT",
            "WT",
            "RT"
        ]

        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header

        for process in self.completed_processes:
            cells = table.add_row().cells

            values = [
                process.pid,
                process.arrival_time,
                process.burst_time,
                process.completion_time,
                process.turnaround_time,
                process.waiting_time,
                process.response_time
            ]

            for index, value in enumerate(values):
                cells[index].text = str(value)

        # Gantt
        document.add_heading(
            "Gantt Chart",
            level=2
        )

        gantt_text = " | ".join(
            f"{pid} ({start}-{end})"
            for pid, start, end in self.gantt_data
        )

        document.add_paragraph(
            gantt_text
        )

        # Average metrics
        count = len(self.completed_processes)

        avg_waiting = sum(
            process.waiting_time
            for process in self.completed_processes
        ) / count

        avg_turnaround = sum(
            process.turnaround_time
            for process in self.completed_processes
        ) / count

        avg_response = sum(
            process.response_time
            for process in self.completed_processes
        ) / count

        document.add_heading(
            "Average Times",
            level=2
        )

        document.add_paragraph(
            f"Average Waiting Time: {avg_waiting:.2f}"
        )

        document.add_paragraph(
            f"Average Turnaround Time: {avg_turnaround:.2f}"
        )

        document.add_paragraph(
            f"Average Response Time: {avg_response:.2f}"
        )

        document.save(file_path)

        messagebox.showinfo(
            "Export DOCX",
            "SJF Preemptive DOCX report exported successfully."
        )